import csv
import logging
from io import TextIOWrapper
from pathlib import Path
from typing import Generator, List

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from crucible.calendar.exceptions import MissingRequiredFields

logging.getLogger('crucible').addHandler(logging.NullHandler())


def fields_check(fields: List) -> bool:
    """
    Checks the fields in the csv file against the list of required fields.

    @param fields:  A list of fields found in the first row of the csv file
    @return:        True or raises a MissingRequiredFields error
    """
    existing_fields = set(fields)
    fields_should_be = {"AudienceLevel", "ClassName", "Description", "Department", "Section", "Hours",
                        "FrequencyRegularPrice", "Prerequisite"}
    if not existing_fields == fields_should_be:
        missing_fields = sorted(fields_should_be.difference(existing_fields))
        msg = f'Invalid fields.  Missing [{", ".join(missing_fields)}]'
        logging.critical(msg)
        raise MissingRequiredFields(msg)
    return True


def create_styled_doc(data: Generator, path: Path) -> None:
    """
    Write the contents of the data generator to path as a styled Word doc.
    """
    logging.debug('Creating styled document')

    document = Document()
    styles = {}

    # Get the first row
    fields = next(data)

    fields_check(fields)

    logging.info('Found styles: %s', ', '.join(fields))

    for field in fields:
        logging.info('Adding style %s', field)
        styles[field] = document.styles.add_style(field, WD_STYLE_TYPE.PARAGRAPH)

    old_section = None
    old_department = None

    for row in data:
        item = dict(zip(fields, row))
        level = item.get('AudienceLevel').strip()
        name = item.get('ClassName').strip()
        description = item.get('Description').strip()
        department = item.get('Department').strip()
        section = item.get('Section').strip()
        hours = item.get('Hours').strip()
        prereq = item.get('Prerequisite', '# PREREQ #').strip()
        meta = item.get('FrequencyRegularPrice').strip()

        logging.info('%s - %s', department, name)

        if old_section != section:
            p = document.add_paragraph(section)
            p.style = styles['Section']

        if old_department != department:
            p = document.add_paragraph(department)
            p.style = styles['Department']

        if name:
            p = document.add_paragraph(name)
            p.style = styles['ClassName']

        if level:
            p = document.add_paragraph(level)
            p.style = styles['AudienceLevel']

        if description:
            p = document.add_paragraph(description)
            p.style = styles['Description']

        if prereq:
            p = document.add_paragraph(prereq)
            p.style = styles['Prerequisite']

        if hours:
            p = document.add_paragraph(hours)
            p.style = styles['Hours']

        if meta:
            p = document.add_paragraph(meta)
            p.style = styles['FrequencyRegularPrice']

        old_section = section
        old_department = department

    document.save(path)

    logging.debug('Finished creating styled document')


def parse_input(fp: TextIOWrapper) -> Generator:
    """
    Reads the input file and returns a generator.
    """
    logging.debug('Parsing input from file')

    data = csv.reader(fp.read().splitlines())
    fp.close()

    logging.debug('Finished parsing input')
    return data


def generate_calendar(args):
    data = parse_input(args.infile)
    create_styled_doc(data=data, path=args.outfile)
    return args.outfile
